from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook


REPO_ROOT = Path(__file__).resolve().parents[2]
PDF_WORKFLOW_SCRIPT = REPO_ROOT / "skills" / "productivity" / "ocr-and-documents" / "scripts" / "build_pdf_workflow_map.py"


def test_docx_create_and_reopen_roundtrip(tmp_path: Path):
    output = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Hermes Verify", level=1)
    doc.add_paragraph("alpha")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "k"
    table.rows[0].cells[1].text = "v"
    table.rows[1].cells[0].text = "x"
    table.rows[1].cells[1].text = "1"
    doc.save(output)

    reopened = Document(output)
    assert output.exists()
    assert output.stat().st_size > 0
    assert len(reopened.paragraphs) >= 2
    assert len(reopened.tables) == 1
    assert next((p.text for p in reopened.paragraphs if p.text.strip()), "") == "Hermes Verify"


def test_xlsx_create_and_reopen_roundtrip(tmp_path: Path):
    output = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "metric"
    ws["B1"] = "value"
    ws["A2"] = "count"
    ws["B2"] = 7
    wb.save(output)

    reopened = load_workbook(output, data_only=False)
    assert output.exists()
    assert output.stat().st_size > 0
    assert reopened.sheetnames == ["Data"]
    ws2 = reopened["Data"]
    assert ws2["A1"].value == "metric"
    assert ws2["B2"].value == 7


def test_pdf_workflow_map_script_outputs_expected_shape():
    result = subprocess.run(
        [sys.executable, str(PDF_WORKFLOW_SCRIPT)],
        capture_output=True,
        text=True,
        check=False,
        timeout=20,
    )
    assert result.returncode == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["entrypoint_count"] >= 4
    assert payload["gap_count"] >= 4

    output = Path(payload["output"])
    data = json.loads(output.read_text(encoding="utf-8"))
    assert data["purpose"] == "Hermes PDF workflow map"
    assert any(item["tool_or_skill"] == "nano-pdf" for item in data["entrypoints"])
    assert "watermarking" in data["gaps"]
